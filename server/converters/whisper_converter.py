import os
import whisper
import docx
import fitz
import torch
from .base import BaseConverter

class WhisperConverter(BaseConverter):
    def __init__(self, source_ext, target_ext, model_size="base"):
        self._source_ext = source_ext
        self._target_ext = target_ext
        self._model_size = model_size
        self._model = None

    @property
    def supported_extension(self):
        return self._source_ext

    @property
    def output_extension(self):
        return self._target_ext

    def _get_model(self, device_pref="auto"):
        target_device = "cpu"
        if device_pref in ["auto", "gpu"]:
            if torch.cuda.is_available():
                target_device = "cuda"
            else:
                try:
                    import torch_directml
                    if torch_directml.is_available():
                        target_device = torch_directml.device()
                except Exception as e:
                    print(f"DirectML initialization failed: {e}")

        need_reload = False
        if self._model is None:
            need_reload = True
        else:
            current_device_str = str(next(self._model.parameters()).device)
            target_device_str = str(target_device)
            if target_device_str == "cuda" and "cuda" not in current_device_str:
                need_reload = True
            elif target_device_str == "cpu" and "cpu" not in current_device_str:
                need_reload = True
            elif "dml" in target_device_str and "dml" not in current_device_str:
                need_reload = True
            elif "privateuseone" in target_device_str and "privateuseone" not in current_device_str:
                need_reload = True

        if need_reload:
            print(f"Loading Whisper model: {self._model_size} on {target_device}...")
            self._model = whisper.load_model(self._model_size, device=target_device)
        return self._model

    def convert(self, file_path, **kwargs):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            device_pref = kwargs.get('device_pref', 'auto')
            model = self._get_model(device_pref=device_pref)
            
            # 执行识别
            # 自动检测是否使用 GPU 以决定 fp16 参数
            use_fp16 = False
            device_str = str(next(model.parameters()).device)
            if "cuda" in device_str:
                use_fp16 = True
            
            try:
                # initial_prompt 引导模型输出简体中文
                result = model.transcribe(
                    file_path, 
                    fp16=use_fp16, 
                    language="zh",
                    initial_prompt="以下是普通话的简体中文。"
                )
            except Exception as e:
                if "dml" in device_str or "privateuseone" in device_str:
                    print(f"GPU (DirectML) transcription failed: {e}. Falling back to CPU...")
                    self._model = whisper.load_model(self._model_size, device="cpu")
                    model = self._model
                    result = model.transcribe(
                        file_path, 
                        fp16=False, 
                        language="zh",
                        initial_prompt="以下是普通话的简体中文。"
                    )
                else:
                    raise e
            text = result["text"].strip()
            
            # 根据目标后缀返回内容
            if self._target_ext == ".md":
                filename = os.path.basename(file_path)
                return f"# Transcription of {filename}\n\n{text}", self._target_ext
            
            elif self._target_ext == ".docx":
                doc = docx.Document()
                doc.add_heading(f"Transcription of {os.path.basename(file_path)}", 0)
                doc.add_paragraph(text)
                
                # 保存到流或内存中？BaseConverter 期望内容。
                # app.py 会将内容写入 BytesIO。
                # 所以我可以直接返回 bytes。
                import io
                doc_io = io.BytesIO()
                doc.save(doc_io)
                return doc_io.getvalue(), self._target_ext
                
            elif self._target_ext == ".pdf":
                # 使用 fitz 创建简单的 PDF
                doc = fitz.open()
                page = doc.new_page()
                
                # 定义中文字体路径 (Windows 常用字体)
                font_path = r"C:\Windows\Fonts\msyh.ttc"
                font_name = "msyh"
                
                # 如果字体不存在，尝试使用内置的或其他备选
                if os.path.exists(font_path):
                    page.insert_font(fontname=font_name, fontfile=font_path)
                else:
                    # 备选：如果是在 Linux 或其他系统，可能需要不同路径
                    font_name = "helv" # 回退到默认，但中文会乱码
                
                # 写入标题
                page.insert_text((50, 72), f"Transcription of {os.path.basename(file_path)}", 
                                 fontsize=16, fontname=font_name)
                
                # 写入正文 (使用 textbox 支持换行)
                rect = fitz.Rect(50, 100, 550, 800)
                page.insert_textbox(rect, text, fontname=font_name, fontsize=11)
                
                pdf_bytes = doc.convert_to_pdf()
                doc.close()
                return pdf_bytes, self._target_ext
            
            return text, self._target_ext
            
        except Exception as e:
            raise Exception(f"Whisper transcription failed: {str(e)}")
